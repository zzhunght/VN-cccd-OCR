import json
from docx import Document
import os
from docx.shared import Pt

def replace_string(doc, data, old_str, new_str, data_extra=None):
    for para in doc.paragraphs:
        if data_extra is not None:
            para.text = para.text.replace(old_str, new_str + data[data_extra])
        else:
            para.text = para.text.replace(old_str, new_str)
    print(data_extra)


def so_yeu_ly_lich(path, json_string):
    doc = Document(path)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Time New Roman'
    font.size = Pt(12)
    data = json.loads(json_string)
    birth_day = data['birth_day'].split('/')
    expire_date = data['expire_date'].split('/') # this is expride_date,

    for para in doc.paragraphs:
        replace_string(doc, data,'Số chứng minh……………',  'Số chứng minh:……', 'id')
        replace_string(doc, data,'Họ và tên (chữ in hoa)………………………………………',  'Họ và tên (chữ in hoa).....', 'name')
        replace_string(doc, data, '. Sinh ngày……tháng …..năm …………', '. Sinh ngày '+
                       birth_day[0] +' tháng '+ birth_day[1] + ' năm '+ birth_day[2]+'   ')
        replace_string(doc, data, 'Nam/ Nữ ….…….……........', 'Nam/ Nữ …..', 'gender')
        replace_string(doc, data, 'Nguyên quán……', 'Nguyên quán…', 'origin_location')
        replace_string(doc, data, 'Chỗ ở hiện nay ……', 'Chỗ ở hiện nay...', 'recent_location')
        #test expire_data == issued_data
        replace_string(doc, data, 'cấp ngày .…/…./……', 'cấp ngày '+ expire_date[0]+'/'+expire_date[1]+'/'+ expire_date[2])
        # replace_string(doc, data, 'cấp ngày .…/…./……', 'cấp ngày '+ issued_date[0]+'/'+issued_date[1]+'/'+ issued_date[2])
        doc.save(f'so-yeu-ly-lich-{data["name"]}.docx')


json_string = ('{"type": "Mặt sau", "name": "PHAM DUY LONG", '
               '"recent_location": "S Trà Co, Thanh Cái, '
               'Qung NInh phó Móng Khu Trang Ginl Trà Co, Thanh Móng Cál, phó", '
               '"birth_day": "03/12/2006", "id": "022206004066", '
               '"expire_date": "01/12/2031", '
               '"nationality": "Việt Nam", '
               '"origin_location": "Hải Xuan, Thành phố Móng Cái, Quảng Ninh Hải Xuán, Thành phó Móng Cá", '
               '"gender": "Nam"}')
path = r'word_form/so-yeu-ly-lich.docx'
so_yeu_ly_lich(path, json_string)

