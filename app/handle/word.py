import json
from docx import Document
import os
from docx.shared import Pt, RGBColor

from fastapi import HTTPException
def check_exists(path):
    exists = os.path.exists(path)
    if exists:
        print('true')
        return True
    else:
        print('false')
        return False

def so_yeu_ly_lich(front_string, back_string):
    # doc = Document(r'D:\learn\VN-cccd-OCR\app\handle\word_form\so-yeu-ly-lich.docx')
    try:
        doc = Document('handle/word_form/so-yeu-ly-lich.docx')
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Time New Roman'
        font.size = Pt(12)
        data = json.loads(front_string)
        dataB = json.loads(back_string)
        birth_day = data['birth_day'].split('/')


        for para in doc.paragraphs:
            for run in para.runs:
                if run.bold is not True:
                    para.text = para.text.replace('Số chứng minh…………………..', 'Số chứng minh……' + data['id']+ '.....')

                    para.text = para.text.replace('Họ và tên (chữ in hoa)………………………………………', 'Họ và tên (chữ in hoa).....'+ data['name']+ '...')
                    para.text = para.text.replace('. Sinh ngày……tháng …..năm …………', '. Sinh ngày ' +
                                    birth_day[0] + ' tháng ' + birth_day[1] + ' năm ' + birth_day[2] + '   ')
                    para.text = para.text.replace('Nam/ Nữ ….…….……........', 'Nam/ Nữ …..' + data['gender'])
                    para.text = para.text.replace('Nguyên quán……', 'Nguyên quán…' + data['origin_location'])
                    para.text = para.text.replace('Chỗ ở hiện nay ……', 'Chỗ ở hiện nay...' + data['recent_location'])
                    
                    para.text = para.text.replace('nơi cấp………', 'nơi cấp…....' + dataB['issue_place'])
                    # para.text = para.text.replace('cấp ngày .…/…./……',
                    #                'cấp ngày ' + data['expire_date'])
                    para.text = para.text.replace('cấp ngày .…/…./……', 'cấp ngày '+ dataB['issue_date'])
        # doc.save(r'D:\learn\VN-cccd-OCR\app\handle\dir_save\so-yeu-ly-lich.docx')
        doc.save(r'handle\dir_save\so-yeu-ly-lich.docx')
        print('saved')
    except Exception as e:
        raise HTTPException(status_code=400, detail='Trích xuất thông tin không thành công')

def don_xin_dk_tam_tru(front_string, back_string):
    # doc = Document(r'D:\learn\VN-cccd-OCR\app\handle\word_form\don_xin_dk_tam_tru.docx')
    try:
        doc = Document('handle/word_form/don_xin_dk_tam_tru.docx')
        data = json.loads(front_string)
        dataB = json.loads(back_string)


        for para in doc.paragraphs:
            for run in para.runs:
                if run.bold is not True:
                    para.text = para.text.replace('Số CCCD: ................................. ', 'Số CMND: ..' + data['id']+ '...')
                    para.text = para.text.replace('Tôi tên là: ....', 'Tôi tên là: ..' + data['name'])
                    para.text = para.text.replace('Ngày sinh:....', 'Ngày sinh:..' +
                                                data['birth_day'] + '   ')
                    para.text = para.text.replace('Địa chỉ thường trú......',
                                                'Địa chỉ thường trú......' + data['recent_location'])
                    # test expire_data == issued_data
                    para.text = para.text.replace('Ngày:..................................................',
                                                'Ngày:...' + data['expire_date'])
                    para.text = para.text.replace('Cấp tại:.....', 'Cấp tại:...' + dataB['issue_place'])
                    para.text = para.text.replace('Ngày cấp:..................................................', 'Ngày cấp:...'+ dataB['issue_date'])
        # doc.save(f'app/app\handle\dir_save\don_xin_dk_tam_tru-{data["name"]}.docx')
        # doc.save(r'D:\learn\VN-cccd-OCR\app\handle\dir_save\don_xin_dk_tam_tru.docx')
        doc.save('handle\dir_save\don_xin_dk_tam_tru.docx')
        print("saved")
    except Exception as e:
        raise HTTPException(status_code=400, detail='Trích xuất thông tin không thành công')

def don_xin_tam_vang(front_string, back_string):
    try:
        # doc = Document(r'D:\learn\VN-cccd-OCR\app\handle\word_form\don_xin_tam_vang.docx')
        doc = Document('handle/word_form/don_xin_tam_vang.docx')
        # D:\ekyc\app\handle\word_form\don_xin_tam_vang.docx
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Time New Roman'
        font.size = Pt(12)
        data = json.loads(front_string)
        dataB = json.loads(back_string)

        for para in doc.paragraphs:
            for para in doc.paragraphs:
                for run in para.runs:
                    if run.bold is not True:
                        para.text = para.text.replace('Số CCCD:................',
                                                    'Số CCCD: ....' + data['id'])
                        para.text = para.text.replace('Tôi tên là: ....', 'Tôi tên là: ..' + data['name'])
                        para.text = para.text.replace('Ngày sinh:....', 'Ngày sinh:..' +
                                                    data['birth_day'] + '   ')
                        para.text = para.text.replace('Địa chỉ thường trú:.....',
                                                    'Địa chỉ thường trú:...' + data['recent_location'])
                        # test expire_data == issued_data
                        para.text = para.text.replace('Ngày:.....',
                                                    'Ngày:...' + data['expire_date'])
                        para.text = para.text.replace('Cấp tại:............................. ', 'Cấp tại:...'+ dataB['issue_place'] + "..")
                        para.text = para.text.replace('Ngày:.....', 'Ngày:...'+ dataB['issue_date'])
        # doc.save(r'D:\learn\VN-cccd-OCR\app\handle\dir_save\don_xin_tam_vang.docx')
        doc.save('handle\dir_save\don_xin_tam_vang.docx')
        print("save")
    except Exception as e:
        raise HTTPException(status_code=400, detail='Trích xuất thông tin không thành công')


json_string2 = json.dumps({'type': 'Căn cược công dân mặt trước',
                           'recent_location': '8, Văn Thân 103/28, Văn Thân 103/28 Phường 08, Quận 6, TP. Hồ Chí Minh',
                           'name': 'PHAN NGỌC TRÂM',
                           'id': '079192019848',
                           'expire_date': '06/06/2032', 'birth_day': '06/06/1992', 'origin_location': 'Phường 08, Quận 6, TP.Hồ Chí Minh Phường 08, Quận 6, TP', 'nationality': 'Việt Nam', 'gender': 'Nữ'})

back = json.dumps({
            'type': 'Căn cược công dân mặt sau',
            'issue_place': 'CỤC TRƯỞNG CỤC CẢNH SÁT QUẢN LÝ HÀNH CHÍNH VỀ TRẬT TỰ XÃ HỘI',
            'issue_date': '13/05/2021'
})
# ======PATH=====

# ======Call Function=====
# so_yeu_ly_lich(json_string2, back)  
# don_xin_dk_tam_tru(json_string2, back)
check_exists("handle\word_form\don_xin_tam_vang.docx")
# don_xin_tam_vang(json_string2, back)

